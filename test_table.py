from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()
doc.add_paragraph("Logo here")

date_table = doc.add_table(rows=1, cols=1)
date_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
date_table.style = 'Table Grid'
date_table.autofit = False

tblPr = date_table._tbl.tblPr
tblW = tblPr.find(qn('w:tblW'))
if tblW is None:
    tblW = OxmlElement('w:tblW')
    tblPr.append(tblW)
tblW.set(qn('w:w'), str(int(2.5 * 1440)))
tblW.set(qn('w:type'), 'dxa')

p_date = date_table.cell(0, 0).paragraphs[0]
p_date.add_run("29.04.2026")

doc.save("test_table.docx")
print("Saved test_table.docx")
