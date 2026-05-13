import io
import os
import qrcode
from barcode import Code128
from barcode.writer import ImageWriter
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_qr(data):
    """Генерує QR-код у вигляді об'єкта BytesIO."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf

def generate_barcode(data):
    """Генерує штрихкод Code128 у вигляді об'єкта BytesIO."""
    # Code128 генерує штрихкод, ImageWriter записує його як зображення
    rv = io.BytesIO()
    Code128(data, writer=ImageWriter()).write(rv)
    rv.seek(0)
    return rv

def generate_act_docx(assets, admin_name="Кульмач Т.В"):
    """Генерує Акт прийому-передачі у форматі DOCX."""
    if isinstance(assets, dict):
        assets = [assets]
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    if os.path.exists("logo.png"):
        doc.add_picture("logo.png", width=Inches(2.0))
        
    date_str = datetime.now().strftime("%d.%m.%Y")
    date_table = doc.add_table(rows=1, cols=1)
    date_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    date_table.style = 'Table Grid'
    date_table.autofit = False
    
    tblPr = date_table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(2.5 * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    
    p_date = date_table.cell(0, 0).paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"{date_str}")
    run_date.font.size = Pt(10)
    
    doc.add_paragraph()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Акт\nПрийому-передачі обладнання")
    run_title.bold = True
    run_title.font.size = Pt(14)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1 + len(assets), cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, ['№', 'Найменування\n(марка, модель)', 'Кількість, шт.', 'Серійний №', 'Примітка']):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for idx, asset in enumerate(assets):
        row_cells = table.rows[idx + 1].cells
        row_cells[0].text = f'{idx + 1}.'
        row_cells[1].text = f"{asset.get('category', '')} {asset.get('brand', '')} {asset.get('model', '')}"
        row_cells[2].text = '1'
        row_cells[3].text = str(asset.get('serial_number', ''))
        row_cells[4].text = str(asset.get('description', ''))
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    
    cell_left = sig_table.cell(0, 0)
    cell_right = sig_table.cell(0, 1)
    
    p_left = cell_left.paragraphs[0]
    p_left.add_run("Майно Передав:\n\n").bold = True
    p_left.add_run(f"ТОВАРИСТВО З ОБМЕЖЕНОЮ\nВІДПОВІДАЛЬНІСТЮ \"ЕЖІС Україна\"\nТОВ \"ЕЖІС Україна\"\nМісце знаходження: вул. Антоновича, буд 29, м. Київ, 01033\n\n\n________________________\nПідпис              {admin_name}")
    
    p_right = cell_right.paragraphs[0]
    p_right.add_run("Майно Прийняв:\n\n\n").bold = True
    assigned_to = str(assets[0].get('assigned_to', ''))
    if assigned_to == 'Не призначено':
        assigned_to = ''
    p_right.add_run(f"________________________\nПідпис              {assigned_to}")
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def generate_return_act_docx(assets, old_assigned_to, admin_name="Кульмач Т.В"):
    """Генерує Акт повернення обладнання у форматі DOCX."""
    if isinstance(assets, dict):
        assets = [assets]
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    if os.path.exists("logo.png"):
        doc.add_picture("logo.png", width=Inches(2.0))
        
    date_str = datetime.now().strftime("%d.%m.%Y")
    date_table = doc.add_table(rows=1, cols=1)
    date_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    date_table.style = 'Table Grid'
    date_table.autofit = False
    
    tblPr = date_table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(2.5 * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    
    p_date = date_table.cell(0, 0).paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"{date_str}")
    run_date.font.size = Pt(10)
    
    doc.add_paragraph()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Акт\nПовернення обладнання")
    run_title.bold = True
    run_title.font.size = Pt(14)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1 + len(assets), cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, ['№', 'Найменування\n(марка, модель)', 'Кількість, шт.', 'Серійний №', 'Примітка']):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for idx, asset in enumerate(assets):
        row_cells = table.rows[idx + 1].cells
        row_cells[0].text = f'{idx + 1}.'
        row_cells[1].text = f"{asset.get('category', '')} {asset.get('brand', '')} {asset.get('model', '')}"
        row_cells[2].text = '1'
        row_cells[3].text = str(asset.get('serial_number', ''))
        row_cells[4].text = str(asset.get('description', ''))
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    
    cell_left = sig_table.cell(0, 0)
    cell_right = sig_table.cell(0, 1)
    
    p_left = cell_left.paragraphs[0]
    p_left.add_run("Майно Передав:\n\n\n").bold = True
    p_left.add_run(f"________________________\nПідпис              {old_assigned_to}")
    
    p_right = cell_right.paragraphs[0]
    p_right.add_run("Майно Прийняв:\n\n").bold = True
    p_right.add_run(f"ТОВАРИСТВО З ОБМЕЖЕНОЮ\nВІДПОВІДАЛЬНІСТЮ \"ЕЖІС Україна\"\nТОВ \"ЕЖІС Україна\"\nМісце знаходження: вул. Антоновича, буд 29, м. Київ, 01033\n\n\n________________________\nПідпис              {admin_name}")
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
def generate_codes_docx(assets, type_choice="QR-код"):
    """Генерує DOCX файл із сіткою QR-кодів або штрихкодів."""
    doc = Document()
    
    # Налаштування вузьких полів
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    table = doc.add_table(rows=(len(assets) + 2) // 3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for idx, asset in enumerate(assets):
        row = idx // 3
        col = idx % 3
        cell = table.cell(row, col)
        
        # Очищуємо клітинку та додаємо контент
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(f"{asset.get('category', '')} {asset.get('brand', '')}\n")
        run.bold = True
        p.add_run(f"SN: {asset.get('serial_number', '')}\n")
        
        # Генерація коду
        if type_choice == "QR-код":
            img_buf = generate_qr(asset['barcode_data'])
            img_width = Inches(1.5)
        else:
            img_buf = generate_barcode(asset['barcode_data'])
            img_width = Inches(2.0)
            
        p.add_run().add_picture(img_buf, width=img_width)
        
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
